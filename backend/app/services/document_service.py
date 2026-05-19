from __future__ import annotations

import base64
import os
from io import BytesIO

import fitz
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader

from app.config import get_settings
from app.services.ai_service import extract_text, get_client
from app.services.storage_service import save_uploaded_file

settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".png", ".jpg", ".jpeg"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
HTML_EXTENSIONS = {".html", ".htm"}


def get_extension(filename: str | None) -> str:
    if not filename:
        raise HTTPException(status_code=400, detail="У файла должно быть имя.")

    extension = os.path.splitext(filename)[1].lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                "Неподдерживаемый формат файла. Разрешены только: "
                ".pdf, .docx, .html, .htm, .png, .jpg, .jpeg."
            ),
        )
    return extension


def ensure_size_limit(file_bytes: bytes) -> None:
    if len(file_bytes) > settings.max_file_size_bytes:
        raise HTTPException(
            status_code=413,
            detail="Файл слишком большой. Максимальный размер файла: 15 MB.",
        )


def parse_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            parts.append(value)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def parse_html(file_bytes: bytes) -> str:
    html_text = file_bytes.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html_text, "html.parser")

    for tag_name in ("script", "style", "noscript"):
        for tag in soup.find_all(tag_name):
            tag.decompose()

    text_blocks: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th"]):
        text = tag.get_text(" ", strip=True)
        if text:
            text_blocks.append(text)

    if text_blocks:
        return "\n".join(text_blocks).strip()

    return soup.get_text("\n", strip=True)


def parse_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages: list[str] = []

    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)

    return "\n\n".join(pages).strip()


def image_bytes_to_data_url(file_bytes: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"


def ocr_image_with_openai(
    file_bytes: bytes,
    mime_type: str,
    filename: str,
) -> str:
    client = get_client()
    response = client.responses.create(
        model=settings.openai_model,
        input=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": (
                            "Извлеки весь читаемый текст с изображения документа. "
                            "Сохрани номера, даты, реквизиты, подписи, заголовки и таблицы. "
                            "Не анализируй, не сокращай, верни только распознанный текст. "
                            f"Имя файла: {filename}."
                        ),
                    },
                    {
                        "type": "input_image",
                        "image_url": image_bytes_to_data_url(file_bytes, mime_type),
                    },
                ],
            }
        ],
    )
    return extract_text(response).strip()


def ocr_scanned_pdf_with_openai(file_bytes: bytes, filename: str) -> str:
    document = fitz.open(stream=file_bytes, filetype="pdf")
    page_texts: list[str] = []

    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image_bytes = pixmap.tobytes("png")
        page_text = ocr_image_with_openai(
            file_bytes=image_bytes,
            mime_type="image/png",
            filename=f"{filename} page {page_index + 1}",
        )
        if page_text:
            page_texts.append(f"[Страница {page_index + 1}]\n{page_text}")

    return "\n\n".join(page_texts).strip()


def parse_upload_file(
    *,
    filename: str,
    content_type: str | None,
    file_bytes: bytes,
) -> tuple[str, str, str]:
    extension = get_extension(filename)

    if extension == ".docx":
        return parse_docx(file_bytes), "docx-text", extension

    if extension in HTML_EXTENSIONS:
        return parse_html(file_bytes), "html-text", extension

    if extension == ".pdf":
        pdf_text = parse_pdf_text(file_bytes)
        if pdf_text:
            return pdf_text, "pdf-text", extension
        return ocr_scanned_pdf_with_openai(file_bytes, filename), "pdf-ocr", extension

    if extension in IMAGE_EXTENSIONS:
        mime_type = content_type or f"image/{extension.lstrip('.')}"
        return ocr_image_with_openai(file_bytes, mime_type, filename), "image-ocr", extension

    raise HTTPException(status_code=400, detail="Формат файла пока не поддерживается.")


def persist_uploaded_file(filename: str, file_bytes: bytes) -> str:
    return save_uploaded_file(filename, file_bytes)
